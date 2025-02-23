import streamlit as st
import lamini
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import os
import re

# 🎨 Custom CSS for Stylish UI
st.markdown("""
    <style>
        body { background-color: #121212; color: #ffffff; font-family: 'Arial', sans-serif; }
        
        /* Input fields */
        .stTextInput>div>div>input, .stTextArea>div>textarea {
            background-color: #f5f5f5 !important;
            color: #333333 !important;
            border-radius: 8px !important;
            padding: 10px !important;
        }
        
        /* Placeholder text */
        ::placeholder {
            color: #666 !important;
            opacity: 1 !important;
        }

        /* Button styling */
        .stButton>button {
            background-color: #007bff !important;
            color: white !important;
            border-radius: 8px !important;
            padding: 8px 16px !important;
            font-size: 16px !important;
        }
        .stButton>button:hover {
            background-color: #0056b3 !important;
        }

        /* Sidebar */
        .sidebar .sidebar-content {
            background-color: #1e1e1e;
            color: white;
        }
    </style>
    """, unsafe_allow_html=True)


# 📌 Lamini LLM Setup
LLM_API_KEY = "YOUR_LAMINI_API_KEY"
llm = lamini.Lamini(api_key=LLM_API_KEY, model_name="llama-2-7b-chat")  # Adjust model name if needed

# 🏠 Sidebar Navigation
st.sidebar.title("⚖️ Legal Assistant")
page = st.sidebar.radio("📌 Choose a Section:", ["💬 Legal Chatbot", "🔍 Case Finder", "📄 Document Generator"])

# 💬 Legal Chatbot Section
if page == "💬 Legal Chatbot":
    st.title("🤖 Legal Chatbot")
    st.write("Ask me anything about IPC, CrPC, or the Constitution!")

    if "messages" not in st.session_state:
        st.session_state["messages"] = []

    for message in st.session_state["messages"]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    user_input = st.chat_input("Type your legal query here...")

    if user_input:
        st.session_state["messages"].append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        response = llm.generate(user_input)

        st.session_state["messages"].append({"role": "assistant", "content": response})
        with st.chat_message("assistant"):
            st.markdown(response)

# 🔍 Case Finder Section
elif page == "🔍 Case Finder":
    st.title("🔍 Find Legal Cases")

    @st.cache_data
    def load_cases():
        return pd.read_csv("case_summaries.csv").dropna()

    df = load_cases()
    case_type = st.text_input("Case Type:")
    case_category = st.text_input("Case Category:")
    legal_provisions = st.text_area("Legal Provisions & Charges:")
    judgment = st.text_area("Judgment & Legal Reasoning:")
    claims = st.text_area("Claims, Settlements & Fines:")

    def find_best_matches(user_input, df, top_n=3):
        df['Combined_Text'] = df[['Case Type', 'Case Category', 'Legal Provisions & Charges', 
                                  'Judgment & Legal Reasoning', 'Claims, Settlements & Fines']].agg(lambda x: ' '.join(x.dropna()), axis=1)
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform(df['Combined_Text'])
        user_vector = vectorizer.transform([user_input])
        similarities = cosine_similarity(user_vector, tfidf_matrix).flatten()
        top_indices = similarities.argsort()[-top_n:][::-1]
        return [(df.iloc[i]['Case Title'], round(similarities[i], 4)) for i in top_indices] if top_indices.size else [("No matching case found.", 0)]

    if st.button("🔍 Find Best Matches"):
        user_input = f"{case_type} {case_category} {legal_provisions} {judgment} {claims}"
        if not user_input.strip():
            st.warning("⚠️ Please enter some details before searching.")
        else:
            best_matches = find_best_matches(user_input, df)
            st.subheader("📌 Top 3 Matching Cases")
            for idx, (case, score) in enumerate(best_matches, 1):
                st.write(f"**{idx}. {case}** (Similarity Score: {score})")

# 📄 Document Generator Section
elif page == "📄 Document Generator":
    st.title("📄 Generate Legal Documents")

    def load_templates():
        return {file.replace("_notice_template.docx", " Notice").replace("_", " ").title(): os.path.join("templates", file)
                for file in os.listdir("templates") if file.endswith("_notice_template.docx")}

    def extract_placeholders(doc_path):
        doc = Document(doc_path)
        return list({match for para in doc.paragraphs for match in re.findall(r"\[(.*?)\]", para.text)})

    def replace_placeholders(input_path, output_path, placeholder_values):
        doc = Document(input_path)
        for para in doc.paragraphs:
            for placeholder, value in placeholder_values.items():
                para.text = para.text.replace(f"[{placeholder}]", value)
        doc.save(output_path)
        st.success(f"✅ Document saved: {output_path}")
        with open(output_path, "rb") as f:
            st.download_button("⬇️ Download Document", f, file_name=os.path.basename(output_path))

    templates = load_templates()
    selected_template = st.selectbox("📄 Choose a Template:", list(templates.keys()))

    if st.button("📥 Load Template"):
        input_path = templates[selected_template]
        if os.path.exists(input_path):
            placeholders = extract_placeholders(input_path)
            st.session_state.update({"placeholders": placeholders, "input_path": input_path,
                                     "output_path": f"generated_{selected_template.replace(' ', '_').lower()}.docx"})

    if "placeholders" in st.session_state:
        st.subheader("✍️ Fill in the required fields:")
        for placeholder in st.session_state["placeholders"]:
            st.session_state[placeholder] = st.text_input(f"Enter value for '{placeholder}':")

        if st.button("📝 Generate Document"):
            if all(st.session_state[placeholder] for placeholder in st.session_state["placeholders"]):
                replace_placeholders(st.session_state["input_path"], st.session_state["output_path"],
                                     {p: st.session_state[p] for p in st.session_state["placeholders"]})
            else:
                st.error("⚠️ Please fill in all fields before proceeding.")
