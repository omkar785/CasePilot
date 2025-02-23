# import streamlit as st
# import pandas as pd
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity

# # Load case data with caching
# @st.cache_data
# def load_cases():
#     file_path = "case_summaries.csv"
#     df = pd.read_csv(file_path)
#     df = df.dropna()  # Remove rows with missing values
#     return df

# # Find top 3 best matches (sorted by highest similarity)
# def find_best_matches(user_input, df, top_n=3):
#     # Combine relevant columns into a single text field
#     df['Combined_Text'] = df[['Case Type', 'Case Category', 'Legal Provisions & Charges', 
#                               'Judgment & Legal Reasoning', 'Claims, Settlements & Fines']].agg(lambda x: ' '.join(x.dropna()), axis=1)

#     # Convert text to TF-IDF vectors
#     vectorizer = TfidfVectorizer()
#     tfidf_matrix = vectorizer.fit_transform(df['Combined_Text'])
#     user_vector = vectorizer.transform([user_input])

#     # Compute cosine similarity
#     similarities = cosine_similarity(user_vector, tfidf_matrix).flatten()

#     # Get top N matches, sorted in descending order (best match first)
#     top_indices = similarities.argsort()[-top_n:][::-1]  
#     results = [(df.iloc[i]['Case Title'], round(similarities[i], 4)) for i in top_indices]

#     return results if results else [("No matching case found.", 0)]

# # Streamlit UI
# st.title("⚖️ Legal Case Finder Chatbot")
# st.write("Enter case details below, and we will find the **top 3 closest matches**, sorted by relevance.")

# df = load_cases()

# # User Inputs
# case_type = st.text_input("Enter Case Type:")
# case_category = st.text_input("Enter Case Category:")
# legal_provisions = st.text_area("Enter Legal Provisions & Charges:")
# judgment = st.text_area("Enter Judgment & Legal Reasoning:")
# claims = st.text_area("Enter Claims, Settlements & Fines:")

# # Button to find matches
# if st.button("🔍 Find Best Matches"):
#     user_input = f"{case_type} {case_category} {legal_provisions} {judgment} {claims}"
    
#     if not user_input.strip():
#         st.warning("⚠️ Please enter some details before searching.")
#     else:
#         best_matches = find_best_matches(user_input, df)
        
#         # Display results
#         st.subheader("📌 Top 3 Matching Cases (Best Match First)")
#         for idx, (case, score) in enumerate(best_matches, 1):
#             st.write(f"**{idx}. {case}** (Similarity Score: {score})")

# import streamlit as st
# import pandas as pd
# import os
# import re
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# from docx import Document

# # Load legal cases dataset
# @st.cache_data
# def load_cases():
#     file_path = "case_summaries.csv"
#     df = pd.read_csv(file_path)
#     df = df.dropna()  # Remove missing values
#     return df

# # Find best matching legal cases
# def find_best_matches(user_input, df, top_n=3):
#     df['Combined_Text'] = df[['Case Type', 'Case Category', 'Legal Provisions & Charges', 
#                               'Judgment & Legal Reasoning', 'Claims, Settlements & Fines']].agg(lambda x: ' '.join(x.dropna()), axis=1)
#     vectorizer = TfidfVectorizer()
#     tfidf_matrix = vectorizer.fit_transform(df['Combined_Text'])
#     user_vector = vectorizer.transform([user_input])
#     similarities = cosine_similarity(user_vector, tfidf_matrix).flatten()
#     top_indices = similarities.argsort()[-top_n:][::-1]  
#     return [(df.iloc[i]['Case Title'], round(similarities[i], 4)) for i in top_indices]

# # Load document templates
# def load_templates():
#     templates_dir = "templates"
#     return {file.replace("_notice_template.docx", " Notice").replace("_", " ").title(): os.path.join(templates_dir, file)
#             for file in os.listdir(templates_dir) if file.endswith("_notice_template.docx")}

# # Extract placeholders from a document
# def extract_placeholders(doc_path):
#     doc = Document(doc_path)
#     pattern = re.compile(r"\[(.*?)\]")
#     placeholders = set()
#     for para in doc.paragraphs:
#         matches = pattern.findall(para.text)
#         placeholders.update(matches)
#     return list(placeholders)

# # Replace placeholders and generate final document
# def replace_placeholders(input_path, output_path, placeholder_values):
#     doc = Document(input_path)
#     for para in doc.paragraphs:
#         for placeholder, value in placeholder_values.items():
#             if f"[{placeholder}]" in para.text:
#                 para.text = para.text.replace(f"[{placeholder}]", value)
#     doc.save(output_path)
#     st.success(f"Document saved: {output_path}")
#     with open(output_path, "rb") as f:
#         st.download_button("Download Document", f, file_name=os.path.basename(output_path))

# # Streamlit UI
# st.set_page_config(page_title="Legal Case & Document Manager", layout="wide")
# st.title("⚖️ Legal Case & Document Manager")

# # Tabs for Case Finder and Document Generator
# tab1, tab2 = st.tabs(["🔍 Find Legal Cases", "📄 Generate Documents"])

# # Legal Case Finder UI
# df = load_cases()
# with tab1:
#     st.header("Find Best Matching Legal Cases")
#     case_type = st.text_input("Enter Case Type:")
#     case_category = st.text_input("Enter Case Category:")
#     legal_provisions = st.text_area("Enter Legal Provisions & Charges:")
#     judgment = st.text_area("Enter Judgment & Legal Reasoning:")
#     claims = st.text_area("Enter Claims, Settlements & Fines:")
    
#     if st.button("🔍 Find Best Matches"):
#         user_input = f"{case_type} {case_category} {legal_provisions} {judgment} {claims}"
#         if not user_input.strip():
#             st.warning("⚠️ Please enter some details before searching.")
#         else:
#             best_matches = find_best_matches(user_input, df)
#             st.subheader("📌 Top 3 Matching Cases")
#             for idx, (case, score) in enumerate(best_matches, 1):
#                 st.write(f"**{idx}. {case}** (Similarity Score: {score})")

# # Document Generator UI
# with tab2:
#     st.header("Generate Legal Documents")
#     templates = load_templates()
#     selected_template = st.selectbox("Choose a template:", list(templates.keys()))
    
#     if st.button("Load Template"):
#         input_path = templates[selected_template]
#         placeholders = extract_placeholders(input_path)
#         st.session_state.placeholders = placeholders
#         st.session_state.input_path = input_path
#         st.session_state.output_path = f"generated_{selected_template.replace(' ', '_').lower()}.docx"
    
#     if 'placeholders' in st.session_state:
#         st.subheader("Fill in the required fields:")
#         for placeholder in st.session_state.placeholders:
#             st.session_state.placeholder_values[placeholder] = st.text_input(f"Enter value for '{placeholder}':")
#         if st.button("Generate Document"):
#             if all(st.session_state.placeholder_values.values()):
#                 replace_placeholders(st.session_state.input_path, st.session_state.output_path, st.session_state.placeholder_values)
#             else:
#                 st.error("Please fill in all fields before proceeding.")


import streamlit as st
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import os
import re

# Set custom theme matching the reference UI
st.markdown(
    """
    <style>
        body {
            background-color: #f5e7db;
            color: #5b4636;
            font-family: 'Arial', sans-serif;
        }
        .stTextInput, .stTextArea, .stSelectbox, .stButton>button {
            background-color: #faebd7;
            color: #5b4636;
            border-radius: 8px;
        }
        .stButton>button:hover {
            background-color: #e8d1bb;
        }
        .sidebar .sidebar-content {
            background-color: #d7c1ac;
            color: #5b4636;
        }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("⚖️ Legal Case & Document Assistant")

# Sidebar navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Choose a section:", ["Case Finder", "Document Generator"])

# Function to load case data
@st.cache_data
def load_cases():
    file_path = "case_summaries.csv"
    df = pd.read_csv(file_path)
    df = df.dropna()
    return df

# Case Finder Section
if page == "Case Finder":
    st.header("🔍 Find Legal Cases")
    df = load_cases()
    
    case_type = st.text_input("Case Type:")
    case_category = st.text_input("Case Category:")
    legal_provisions = st.text_area("Legal Provisions & Charges:")
    judgment = st.text_area("Judgment & Legal Reasoning:")
    claims = st.text_area("Claims, Settlements & Fines:")
    
    def find_best_matches(user_input, df, top_n=3):
        df['Combined_Text'] = df[['Case Type', 'Case Category', 'Legal Provisions & Charges', 
                                  'Judgment & Legal Reasoning', 'Claims, Settlements & Fines']].agg(lambda x: ' '.join(x.dropna()), axis=1)
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform(df['Combined_Text'])
        user_vector = vectorizer.transform([user_input])
        similarities = cosine_similarity(user_vector, tfidf_matrix).flatten()
        top_indices = similarities.argsort()[-top_n:][::-1]  
        results = [(df.iloc[i]['Case Title'], round(similarities[i], 4)) for i in top_indices]
        return results if results else [("No matching case found.", 0)]
    
    if st.button("Find Best Matches"):
        user_input = f"{case_type} {case_category} {legal_provisions} {judgment} {claims}"
        if not user_input.strip():
            st.warning("⚠️ Please enter some details before searching.")
        else:
            best_matches = find_best_matches(user_input, df)
            st.subheader("📌 Top 3 Matching Cases")
            for idx, (case, score) in enumerate(best_matches, 1):
                st.write(f"**{idx}. {case}** (Similarity Score: {score})")

# Document Generator Section
elif page == "Document Generator":
    st.header("📄 Generate Legal Documents")
    
    def load_templates():
        templates_dir = "templates"
        return {file.replace("_notice_template.docx", " Notice").replace("_", " ").title(): os.path.join(templates_dir, file)
                for file in os.listdir(templates_dir) if file.endswith("_notice_template.docx")}
    
    def extract_placeholders(doc_path):
        doc = Document(doc_path)
        pattern = re.compile(r"\[(.*?)\]")
        placeholders = {match for para in doc.paragraphs for match in pattern.findall(para.text)}
        return list(placeholders)
    
    def replace_placeholders(input_path, output_path, placeholder_values):
        doc = Document(input_path)
        for para in doc.paragraphs:
            for placeholder, value in placeholder_values.items():
                if f"[{placeholder}]" in para.text:
                    para.text = para.text.replace(f"[{placeholder}]", value)
        doc.save(output_path)
        st.success(f"Document saved: {output_path}")
        with open(output_path, "rb") as f:
            st.download_button("Download Document", f, file_name=os.path.basename(output_path))
    
    templates = load_templates()
    selected_template = st.selectbox("Choose a template:", list(templates.keys()))
    
    if st.button("Load Template"):
        input_path = templates[selected_template]
        if os.path.exists(input_path):
            placeholders = extract_placeholders(input_path)
            st.session_state.placeholders = placeholders
            st.session_state.input_path = input_path
            st.session_state.output_path = f"generated_{selected_template.replace(' ', '_').lower()}.docx"
    
    if 'placeholders' in st.session_state:
        st.subheader("Fill in the required fields:")
        for placeholder in st.session_state.placeholders:
            st.session_state[placeholder] = st.text_input(f"Enter value for '{placeholder}':")
        
        if st.button("Generate Document"):
            if all(st.session_state[placeholder] for placeholder in st.session_state.placeholders):
                replace_placeholders(
                    st.session_state.input_path,
                    st.session_state.output_path,
                    {p: st.session_state[p] for p in st.session_state.placeholders}
                )
            else:
                st.error("Please fill in all fields before proceeding.")

