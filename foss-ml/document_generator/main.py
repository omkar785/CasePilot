import streamlit as st
from docx import Document
import os
import re

def load_templates():
    templates_dir = "templates"
    return {file.replace("_notice_template.docx", " Notice").replace("_", " ").title(): os.path.join(templates_dir, file)
            for file in os.listdir(templates_dir) if file.endswith("_notice_template.docx")}

def extract_placeholders(doc_path):
    doc = Document(doc_path)
    placeholders = set()

    # Find all placeholders in paragraphs
    pattern = re.compile(r"\[(.*?)\]")

    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        placeholders.update(matches)

    return list(placeholders)

def replace_placeholders(input_path, output_path, placeholder_values):
    doc = Document(input_path)

    # Replace placeholders with user input
    for para in doc.paragraphs:
        for placeholder, value in placeholder_values.items():
            if f"[{placeholder}]" in para.text:
                para.text = para.text.replace(f"[{placeholder}]", value)

    # Save the modified document
    doc.save(output_path)
    st.success(f"Document saved: {output_path}")

    # Provide download link
    with open(output_path, "rb") as f:
        st.download_button("Download Document", f, file_name=os.path.basename(output_path))

def main():
    st.title("Document Generator")

    # Initialize session state for placeholder values
    if 'placeholder_values' not in st.session_state:
        st.session_state.placeholder_values = {}

    # Upload a new template
    st.sidebar.header("Upload a New Template")
    uploaded_file = st.sidebar.file_uploader("Upload a .docx template", type=["docx"])

    if uploaded_file:
        template_title = st.sidebar.text_input("Enter title for the template (e.g., Custom Notice):")
        if st.sidebar.button("Save Template"):
            if template_title:
                template_path = f"templates/{template_title.replace(' ', '_').lower()}_notice_template.docx"
                with open(template_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                st.sidebar.success(f"Template saved as: {template_path}")
            else:
                st.sidebar.error("Please provide a valid template title.")

    # Template format guide in the sidebar
    st.sidebar.header("Template Format Guide")
    st.sidebar.markdown("- Use placeholders in the format: `[var_name]` where you want user input.")
    st.sidebar.markdown("- Example: 'Date: [date]' will prompt the user to enter a value for 'date'.")

    # Load available templates
    templates = load_templates()

    # User input: Select template
    selected_template = st.selectbox("Choose a template:", list(templates.keys()))

    if st.button("Load Template"):
        input_path = templates[selected_template]

        if os.path.exists(input_path):
            # Extract placeholders
            placeholders = extract_placeholders(input_path)

            # Store placeholders in session state
            st.session_state.placeholders = placeholders
            st.session_state.input_path = input_path
            st.session_state.output_path = f"generated_{selected_template.replace(' ', '_').lower()}.docx"

    # Check if placeholders are loaded
    if 'placeholders' in st.session_state:
        st.subheader("Fill in the required fields:")

        # Collect inputs for each placeholder
        for placeholder in st.session_state.placeholders:
            st.session_state.placeholder_values[placeholder] = st.text_input(
                f"Enter value for '{placeholder}':",
                value=st.session_state.placeholder_values.get(placeholder, "")
            )

        if st.button("Generate Document"):
            # Ensure all inputs are filled before generating the document
            if all(st.session_state.placeholder_values.values()):
                replace_placeholders(
                    st.session_state.input_path,
                    st.session_state.output_path,
                    st.session_state.placeholder_values
                )
            else:
                st.error("Please fill in all fields before proceeding.")

if __name__ == "__main__":
    main()
